import docx
import re
import os
import zipfile
import shutil
import random

# File paths
input_docx = r"C:\Users\himan\Downloads\Red Herring Prospectus.docx"
output_docx = r"C:\Users\himan\Downloads\Red Herring Prospectus_Redacted.docx"
scratch_dir = r"C:\Users\himan\.gemini\antigravity-ide\brain\0d2eaceb-51e9-434a-aa46-9c8e92ee46c1\scratch"

# Define dynamic mappings for regex items
email_map = {}
phone_map = {}
pan_map = {}
aadhaar_map = {}

# Predefined replacements list (from longest to shortest to avoid partial replacement bugs)
predefined_replacements = [
    # Full Names with variations/decorations
    ("Kushal Subbayya Hegde*", "Vikram Kumar Sen*"),
    ("Pushpa Kushal Hegde*", "Priya Vikram Sen*"),
    ("Rajesh Kushal Hegde*", "Rohan Vikram Sen*"),
    ("Rohit Kushal Hegde*", "Rahul Vikram Sen*"),
    ("Rakhi Girija Shetty&", "Ritu Kumari Sharma&"),
    ("Rakhi Girija Shetty^&", "Ritu Kumari Sharma^&"),
    ("Rakhi Girija Shetty&", "Ritu Kumari Sharma&"),
    
    ("Lalit Muljibhai Sarvaiya", "Lokesh Kumar Sharma"),
    ("Katyayani Balasubramanian", "Kavita Subramanian"),
    ("Kushal Subbayya Hegde", "Vikram Kumar Sen"),
    ("Dinesh Hirachand Munot", "Deepak Kumar Mehta"),
    ("Shanti Gopalkrishnan", "Sneha Subramanian"),
    ("Sarthak Malvadkar", "Sandeep More"),
    ("Rajesh Kushal Hegde", "Rohan Vikram Sen"),
    ("Pushpa Kushal Hegde", "Priya Vikram Sen"),
    ("Karunakar Bhandary", "Karunakar Bose"),
    ("Sunil Nagayya Shetty", "Suresh Kumar Sharma"),
    ("Rohit Kushal Hegde", "Rahul Vikram Sen"),
    ("Rakhi Girija Shetty", "Ritu Kumari Sharma"),
    ("Salil Ajay Bhargava", "Saurabh Amit Bhargava"),
    ("Ramesh Kumar Tiwari", "Ramesh Kumar Tripathi"),
    ("Ram Kumar Tiwari", "Ramesh Kumar Tripathi"),
    ("Sandesh Bhagwat", "Sanjay Bhatt"),
    ("Lokesh Shah", "Lalit Patel"),
    ("Soumavo Sarkar", "Siddharth Sen"),
    ("Kishan Rastogi", "Karan Sharma"),
    ("Abhijit Diwan", "Abhishek Deshmukh"),
    ("Maithili Rajesh Hegde", "Meera Rohan Sen"),
    ("Jabeen Ajay Menon", "Jyoti Amit Nair"),
    ("Ajay Menon", "Amit Nair"),
    ("Ganesh Prasad", "Girish Sharma"),
    ("Prakash Boricha", "Pranav Bhatt"),
    ("Sheetal Parab", "Shruti Patel"),
    ("Siddharth Jadhav", "Soham Joshi"),
    ("Sachin Gawade", "Sameer Gokhale"),
    ("Eric Bacha", "Ethan Brown"),
    ("Tushar Gavankar", "Tanmay Patil"),
    ("Pravin Teli", "Pradeep Tambe"),
    ("Varun Badai", "Vijay Birla"),
    ("Cherag Gyara", "Chetan Gandhi"),
    ("Manisha Shukla", "Mansi Sharma"),
    ("Ashish Mathew Pulloor", "Alok Mathew"),
    ("Anand Soni", "Aditya Shah"),
    ("Chitra Raste", "Charu Rao"),
    ("Sharmila Joshi", "Shalini Kulkarni"),
    ("Parag Pansare", "Pankaj Patil"),
    ("Hitesh Ramani", "Harsh Roy"),
    ("Vijay Khantwal", "Vinay Kumar"),
    ("Vivek Kumar Yadav", "Vikas Yadav"),
    ("Vishal Singh", "Vikram Singh"),
    ("Sugriv Singh", "Suresh Singh"),
    ("DM Shetty", "D.M. Sharma"),
    ("Gopal BO", "G. B. Patel"),
    ("SA Shetty", "S.A. Sharma"),
    ("Jayaram Shetty", "Jayaram Sharma"),
    ("Karunakar Hegde", "Karunakar Sen"),
    ("Narayana B. Shetty", "Narayana Sharma"),
    ("Karunakar N. Bhandary", "Karunakar Bose"),
    ("Narayna B. Shetty", "Narayana Sharma"),
    ("Jayaram N. Shetty", "Jayaram Sharma"),
    ("Ajay Shriram Patil", "Amit Shriram Patel"),
    ("Indu Jacob", "Isha Joshi"),
    ("Amod Joshi", "Anil Kulkarni"),
    
    # 2-word Names
    ("Kushal Subbayya", "Vikram Kumar"),
    ("Pushpa Kushal", "Priya Vikram"),
    ("Rajesh Kushal", "Rohan Vikram"),
    ("Rohit Kushal", "Rahul Vikram"),
    ("Rakhi Girija", "Ritu Kumari"),
    ("Dinesh Hirachand", "Deepak Kumar"),
    ("Ajay Shriram", "Amit Shriram"),
    ("Lalit Muljibhai", "Lokesh Kumar"),
    ("Maithili Rajesh", "Meera Rohan"),
    ("Salil Ajay", "Saurabh Amit"),
    ("Jabeen Ajay", "Jyoti Amit"),
    ("Sunil Nagayya", "Suresh Kumar"),
    
    # Addresses
    ("11/3, 11/4 and 11/5 Village Birdewadi, Chakan Taluka - Khed, Pune – 410 501, Maharashtra, India", "Plot 123, Sector 4, Industrial Area, Chakan, Pune – 410501, Maharashtra, India"),
    ("11/3, 11/4 and 11/5 Village Birdewadi Chakan Taluka - Khed Pune – 410 501 Maharashtra, India", "Plot 123, Sector 4, Industrial Area, Chakan, Pune – 410501, Maharashtra, India"),
    ("11/3, 11/4 and 11/5, Village Birdewadi, Chakan Taluka-Khed, Pune – 410 501, Maharashtra, India", "Plot 123, Sector 4, Industrial Area, Chakan, Pune – 410501, Maharashtra, India"),
    ("11/3, 11/4 and 11/5, Village Birdewadi, Chakan Taluka - Khed, Pune – 410 501, Maharashtra, India", "Plot 123, Sector 4, Industrial Area, Chakan, Pune – 410501, Maharashtra, India"),
    ("201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner, Pune – 411 045, Maharashtra, India", "Suite 401, Level 4, Plaza Business Center, Baner Road, Pune – 411045, Maharashtra, India"),
    ("201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner Pune – 411 045 Maharashtra, India", "Suite 401, Level 4, Plaza Business Center, Baner Road, Pune – 411045, Maharashtra, India"),
    ("S. no. 245/ 104, Pushpakamal, Deccan Gymkhana Society, lane no. 3 Prabhat Road, opposite PYC basketball court, Deccan Gymkhana, Pune – 411 004 Maharashtra, India", "Flat 101, Orchid Residency, lane no. 2, Prabhat Road, Pune – 411004, Maharashtra, India"),
    ("12 Buena Monte, NCL co-operative housing society, Panchvati, Pashan, Pune – 411 008, Maharashtra, India", "Flat 302, Green Meadows Society, Pashan, Pune – 411008, Maharashtra, India"),
    ("Pushpakamal Apartment, Flat – 1, S. no. 245/ 104, Prabhat Road Lane no. 3, Shivaji Nagar, Deccan Gymkhana, Pune – 411 004, Maharashtra, India", "Flat 101, Orchid Residency, lane no. 2, Prabhat Road, Pune – 411004, Maharashtra, India"),
    ("Pratik Bunglow, Senapati Bapat Road, behind Sahara Hotel, Shivajinagar, Model Colony, Pune – 411 016, Maharashtra, India", "Bungalow No. 5, Model Colony, Pune – 411016, Maharashtra, India"),
    ("602, Gopalkrupa Apartment, Bhonde colony, Prabhat Road, Erandawane, Pune – 411 004, Maharashtra, India", "Flat 602, Landmark Apartment, Erandawane, Pune – 411004, Maharashtra, India"),
    ("A-259, JK Road, Minal Residency, Huzur, Govindpura, Bhopal – 462 023, Madhya Pradesh, India", "House No. 12, JK Road, Minal Residency, Bhopal – 462023, Madhya Pradesh, India"),
    ("A29, Abhimanshree Society, Pashan Road, Pune – 411 008, Maharashtra, India", "House No. 4, Abhimanshree Society, Pashan Road, Pune – 411008, Maharashtra, India"),
    ("801 - 804, Wing A, Building No 3, Inspire BKC, G Block, Bandra Kurla Complex, Bandra East, Mumbai 400051, Maharashtra, India", "Suite 101, G-Block, BKC, Bandra East, Mumbai 400051, Maharashtra, India"),
    ("ICICI Venture House, Appasaheb Marathe Marg, Prabhadevi, Mumbai 400025, Maharashtra, India", "Financial Center, Appasaheb Marathe Marg, Prabhadevi, Mumbai 400025, Maharashtra, India"),
    ("C-101, Embassy 247, 1st Floor, L B S Marg, Vikhroli (West), Mumbai 400083, (Maharashtra), India", "Tower B, LBS Marg, Vikhroli West, Mumbai 400083, Maharashtra, India"),
    ("Plot No. J-25, Taloja Industrial Area, Village Padghe, Taluka Panvel, Raigad – 410 208, Maharashtra, India", "Plot No. 44, Taloja Industrial Area, Panvel, Raigad – 410208, Maharashtra, India"),
    ("Plot No. 5, Chakan Industrial Area, Phase II, Village Khalumbre, Taluka Khed, Pune – 410 501, Maharashtra, India", "Plot No. 88, Chakan Industrial Area, Phase II, Pune – 410501, Maharashtra, India"),
    ("Plot No. F-223, Supa Parner Industrial Park, Mauje Palve Khurd, Taluka Parner, Dist – Ahmednagar, Maharashtra – 414 301", "Plot No. A-12, Supa Industrial Area, Parner, Ahmednagar, Maharashtra – 414301"),
    ("PCNTDA Green Building Block A 1st and 2nd floor Near Akurdi Railway Station Akurdi, Pune – 411 044 Maharashtra, India", "Pimpri Business Hub, Near Akurdi Railway Station, Akurdi, Pune – 411044, Maharashtra, India"),
    ("5th Floor, Wing A, Gopal House S. No. 127/1B/1, Plot A1 Opp Harshal Hall Kothrud Pune – 411 038 Maharashtra, India", "5th Floor, Gopal Tower, Kothrud, Pune – 411038, Maharashtra, India"),
    ("Flat No. 102, Sai Complex Shaniwar Peth, Pune – 411 030 Maharashtra, India", "Flat 102, Shaniwar Peth, Pune – 411030, Maharashtra, India"),
    ("8th Floor, Onyx Tower North Main Road Koregaon Park, Pune – 411 001 Maharashtra, India", "Onyx Business Tower, Koregaon Park, Pune – 411001, Maharashtra, India"),
    ("No. 401, 401(A), 401(B) & 402, 402(A), 402(B), 4th Floor Signature Building, Bhandarkar road Shivaji Nagar, Pune – 411 004 Maharashtra, India", "4th Floor, Signature Tower, Bhandarkar Road, Pune – 411004, Maharashtra, India"),
    ("362, Satguru House Next to Tanishq Showroom, CTS No. 30 Bund Garden Road, Pune – 411 001 Maharashtra, India", "Suite 30, Bund Garden Road, Pune – 411001, Maharashtra, India"),
    ("5th Floor, Marathon IT Park Bund Garden Road Pune – 411 001 Maharashtra, India", "Marathon Tower, Bund Garden Road, Pune – 411001, Maharashtra, India"),
    ("Tara Chambers, Mumbai-Pune Road, Wakdewadi Pune – 411 003 Maharashtra, India", "Tara Tower, Wakdewadi, Pune – 411003, Maharashtra, India"),
    ("Ground Floor, Kubera Chambers Opp. Sancheti Hospital Shivajinagar, Pune – 411 005 Maharashtra, India", "Kubera House, Shivajinagar, Pune – 411005, Maharashtra, India"),
    ("Unit no. 1601, B- wing BKC, Mumbai Maharashtra India", "Suite 1601, BKC, Mumbai, Maharashtra, India"),
    ("One World Centre 10th Floor, Tower 2A & 2B Senapati Bapat Marg, Lower Parel (West) Mumbai – 400 013 Maharashtra, India", "World Legal Center, Lower Parel, Mumbai – 400013, Maharashtra, India"),
    ("FIG-OPS Department – Lodha I Think Techno Campus, O-3 Level Next to Kanjurmarg Railway Station, Kanjurmarg (East) Mumbai – 400 042, Maharashtra, India", "Techno Campus, Kanjurmarg East, Mumbai – 400042, Maharashtra, India"),

    # Company Names
    ("KSH INTERNATIONAL LIMITED", "ABC MANUFACTURING LIMITED"),
    ("KSH INTERNATIONAL PRIVATE LIMITED", "ABC MANUFACTURING PRIVATE LIMITED"),
    ("KSH International Limited", "ABC Manufacturing Limited"),
    ("KSH International Private Limited", "ABC Manufacturing Private Limited"),
    ("Bhandary Metal Extrusion Private Limited", "XYZ Metal Extrusion Private Limited"),
    ("Waterloo Industrial Park VI Private Limited", "Riverdale Industrial Park VI Private Limited"),
    ("Waterloo Industrial Park I Private Limited", "Riverdale Industrial Park I Private Limited"),
    ("Waterloo Industrial Park II Private Limited", "Riverdale Industrial Park II Private Limited"),
    ("Waterloo Industrial Park III Private Limited", "Riverdale Industrial Park III Private Limited"),
    ("Waterloo Industrial Park IV Private Limited", "Riverdale Industrial Park IV Private Limited"),
    ("Waterloo Industrial Park V Private Limited", "Riverdale Industrial Park V Private Limited"),
    ("Waterloo Industrial Park VIII Private Limited", "Riverdale Industrial Park VIII Private Limited"),
    ("Waterloo Industrial Park IX Private Limited", "Riverdale Industrial Park IX Private Limited"),
    ("Waterloo Industrial Park IX B Private Limited", "Riverdale Industrial Park IX B Private Limited"),
    ("Waterloo Industrial Park IX A Private Limited", "Riverdale Industrial Park IX A Private Limited"),
    ("Waterloo Motors Private Limited", "Riverdale Motors Private Limited"),
    ("KSH Project Management Services Private Limited", "ABC Project Management Services Private Limited"),
    ("KSH Infra Park 5 Private Limited", "ABC Infra Park 5 Private Limited"),
    ("KSH Infra Park VI Private Limited", "ABC Infra Park VI Private Limited"),
    ("KSH Distriparks Private Limited", "ABC Distriparks Private Limited"),
    ("KSH Integrated Logistics Private Limited", "ABC Integrated Logistics Private Limited"),
    ("Kushal Motors and Electricals Private Limited", "Vikram Motors and Electricals Private Limited"),
    ("KSH Infra Park IV Private Limited", "ABC Infra Park IV Private Limited"),
    ("KSH Infra Park 4 Private Limited", "ABC Infra Park 4 Private Limited"),
    ("Nuvama Wealth Management Limited", "Alpha Wealth Management Limited"),
    ("ICICI Securities Limited", "Beta Securities Limited"),
    ("MUFG Intime India Private Limited", "Delta Registrar Services Private Limited"),
    ("Link Intime India Private Limited", "Omega Share Transfer Private Limited"),
    ("HDFC Bank Limited", "Prime Commercial Bank Limited"),
    ("ICICI Bank Limited", "Apex Bank Limited"),
    ("State Bank of India", "National Trust Bank of India"),
    ("The Federal Bank Limited", "Federal Union Bank Limited"),
    ("Bajaj Finance Limited", "Apex Finance Limited"),
    ("Citibank N.A.", "Global Bank N.A."),
    ("Export-Import Bank of India", "Exim Bank of India"),
    ("IndusInd Bank Limited", "Indus Bank Limited"),
    ("Kirtane & Pandit LLP", "K. P. & Associates LLP"),
    ("CARE Analytics and Advisory Private Limited", "Rating Analytics Private Limited"),
    ("CareEdge Research", "Edge Research"),
    ("Parijat Foundation", "Lotus Foundation"),
    ("Kushal Electricals", "Vikram Electricals"),
    ("Waterloo Motors", "Riverdale Motors"),
    ("Shubhkamal Leasing and Investment Private Limited", "Shubh Leasing and Investment Private Limited"),

    # Domains & Short Company Brand names
    ("kshinternational.com", "abcinternational.com"),
    ("kshinterantional.com", "abcinternational.com"),
    ("kshinternational", "abcinternational"),
    ("Waterloo", "Riverdale"),
    ("Nuvama", "Alpha"),
    ("Trilegal", "TriLaw"),
    ("Kirtane & Pandit", "K. P. & Associates"),
    ("Kirtane", "K. P."),
    ("CareEdge", "Edge"),
    ("HDFC", "Apex"),
    ("ICICI", "Beta"),
    ("IndusInd", "Indus"),
    ("Parijat", "Lotus"),
    ("NSDL", "National Depository Services Limited"),
    ("CDSL", "Central Depository Services Limited"),
    ("KSH", "ABC"),

    # Direct ID numbers
    ("U28129PN1979PLC141032", "U99999PN9999PLC999999"), # CIN
    ("M-140388", "M-999999"), # Engineer Reg
    ("INM000013004", "INM999999999"), # SEBI Reg
    ("INM000011179", "INM999999998"), # SEBI Reg
    ("INR000004058", "INR999999999"), # SEBI Reg
    ("INZ0000166136", "INZ999999999"), # SEBI Reg
    
    # DINs
    ("00135070", "09999999"),
    ("00114193", "09999998"),
    ("00134926", "09999997"),
    ("03124510", "09999996"),
    ("00049801", "09999995"),
    ("01217000", "09999994"),
    ("10938958", "09999993"),
    ("05293084", "09999992"),

    # Individual First and Last Names
    ("Kushal", "Vikram"),
    ("Pushpa", "Priya"),
    ("Rajesh", "Rohan"),
    ("Rohit", "Rahul"),
    ("Rakhi", "Ritu"),
    ("Dinesh", "Deepak"),
    ("Ajay", "Amit"),
    ("Sandesh", "Sanjay"),
    ("Amod", "Anil"),
    ("Sarthak", "Sandeep"),
    ("Lalit", "Lokesh"),
    ("Shanti", "Sneha"),
    ("Lokesh", "Lalit"),
    ("Soumavo", "Siddharth"),
    ("Kishan", "Karan"),
    ("Abhijit", "Abhishek"),
    ("Maithili", "Meera"),
    ("Katyayani", "Kavita"),
    ("Salil", "Saurabh"),
    ("Jabeen", "Jyoti"),
    ("Sunil", "Suresh"),
    ("Ganesh", "Girish"),
    ("Prakash", "Pranav"),
    ("Sheetal", "Shruti"),
    ("Siddharth", "Soham"),
    ("Sachin", "Sameer"),
    ("Eric", "Ethan"),
    ("Tushar", "Tanmay"),
    ("Pravin", "Pradeep"),
    ("Varun", "Vijay"),
    ("Cherag", "Chetan"),
    ("Manisha", "Mansi"),
    ("Ashish", "Alok"),
    ("Anand", "Aditya"),
    ("Chitra", "Charu"),
    ("Sharmila", "Shalini"),
    ("Parag", "Pankaj"),
    ("Hitesh", "Harsh"),
    ("Vijay", "Vinay"),
    ("Vivek", "Vikas"),
    ("Vishal", "Vineet"),
    ("Sugriv", "Suresh"),
    
    ("Hegde", "Sen"),
    ("Shetty", "Sharma"),
    ("Bhagwat", "Bhatt"),
    ("Joshi", "Kulkarni"),
    ("Malvadkar", "More"),
    ("Sarvaiya", "Sharma"),
    ("Munot", "Mehta"),
    ("Patil", "Patel"),
    ("Tiwari", "Tripathi"),
    ("Jacob", "Joshi"),
    ("Rastogi", "Sharma"),
    ("Diwan", "Deshmukh"),
    ("Bhargava", "Bhattacharya"),
    ("Menon", "Nair"),
    ("Boricha", "Bhatt"),
    ("Parab", "Patel"),
    ("Jadav", "Joshi"),
    ("Gawade", "Gokhale"),
    ("Bacha", "Brown"),
    ("Gavankar", "Patil"),
    ("Teli", "Tambe"),
    ("Badai", "Birla"),
    ("Gyara", "Gandhi"),
    ("Shukla", "Sharma"),
    ("Wakhele", "Wagle"),
    ("Pulloor", "Mathew"),
    ("Soni", "Shah"),
    ("Raste", "Rao"),
    ("Pansare", "Patil"),
    ("Ramani", "Roy"),
    ("Khantwal", "Kumar"),
    ("Yadav", "Verma"),
    ("Bhandary", "Bose"),
]

# Generate term to placeholder mapping
term_to_placeholder = {}
placeholder_to_fake = {}

for idx, (old_val, new_val) in enumerate(predefined_replacements):
    placeholder = f"##VAL_{idx}##"
    term_to_placeholder[old_val] = placeholder
    placeholder_to_fake[placeholder] = new_val

substring_allowed = {
    "KSH", "Waterloo", "Nuvama", "Trilegal", "CareEdge", "HDFC", "ICICI", "IndusInd", "Parijat",
    "kshinternational", "kshinterantional", "kshinternational.com", "kshinterantional.com"
}

# Compile a single union regex for all predefined terms
sorted_terms = sorted(term_to_placeholder.keys(), key=len, reverse=True)
union_pattern = re.compile("|".join(re.escape(t) for t in sorted_terms))

def apply_regex_placeholders(text):
    # 1. Emails
    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    for email in re.findall(email_pattern, text):
        if email not in email_map:
            parts = email.split('@')
            user = parts[0]
            domain = parts[1]
            fake_user = "".join(random.choice("abcdefghijklmnopqrstuvwxyz") for _ in range(5))
            fake_domain = "example.com"
            if "nuvama" in domain:
                fake_domain = "alphacapital.com"
            elif "icicisecurities" in domain:
                fake_domain = "betasecurities.com"
            elif "mufg" in domain or "mpms" in domain:
                fake_domain = "deltaregistrar.com"
            elif "hdfcbank" in domain:
                fake_domain = "primebank.com"
            elif "icicibank" in domain:
                fake_domain = "apexbank.com"
            elif "trilegal" in domain:
                fake_domain = "trilaw.com"
            elif "kirtanepandit" in domain:
                fake_domain = "kp-audit.com"
            elif "citi" in domain:
                fake_domain = "globalbank.com"
            elif "indusind" in domain:
                fake_domain = "indusbank.com"
            elif "nsdl" in domain:
                fake_domain = "nationaldepository.com"
            email_map[email] = f"{fake_user}@{fake_domain}"
        text = text.replace(email, email_map[email])
        
    # 2. Phone numbers
    phone_pattern = r'(\+91\s?\d{2,4}\s?\d{7,8}|\+91-\d{2}-\d{8}|022-\d{8}|\b\d{10}\b)'
    for phone in re.findall(phone_pattern, text):
        if phone not in phone_map:
            digits = "".join(random.choice("0123456789") for _ in range(8))
            if phone.startswith("+91"):
                fake_phone = f"+91 98{digits}"
            elif phone.startswith("022"):
                fake_phone = f"022-2{digits[:7]}"
            else:
                fake_phone = f"98{digits}"
            phone_map[phone] = fake_phone
        text = text.replace(phone, phone_map[phone])

    # 3. PAN numbers
    pan_pattern = r'\b[A-Z]{5}[0-9]{4}[A-Z]\b'
    for pan in re.findall(pan_pattern, text):
        if pan not in pan_map:
            fake_pan = "ABCDE" + str(random.randint(1000, 9999)) + "Z"
            pan_map[pan] = fake_pan
        text = text.replace(pan, pan_map[pan])

    # 4. Aadhaar numbers
    aadhaar_pattern = r'\b\d{4}\s\d{4}\s\d{4}\b'
    for aadhaar in re.findall(aadhaar_pattern, text):
        if aadhaar not in aadhaar_map:
            fake_aadhaar = f"{random.randint(1000, 9999)} {random.randint(1000, 9999)} {random.randint(1000, 9999)}"
            aadhaar_map[aadhaar] = fake_aadhaar
        text = text.replace(aadhaar, aadhaar_map[aadhaar])
        
    # 5. IP addresses
    ip_pattern = r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b'
    for ip in re.findall(ip_pattern, text):
        text = text.replace(ip, "192.168.1.100")
        
    return text

def fill_placeholders(text):
    for placeholder, fake in placeholder_to_fake.items():
        text = text.replace(placeholder, fake)
    return text

def replace_term_in_paragraph(paragraph, term, placeholder, is_whole_word=False):
    pt = paragraph.text
    matches = []
    if is_whole_word:
        for m in re.finditer(r'\b' + re.escape(term) + r'\b', pt):
            matches.append((m.start(), m.end()))
    else:
        start = 0
        while True:
            idx = pt.find(term, start)
            if idx == -1:
                break
            matches.append((idx, idx + len(term)))
            start = idx + len(term)
            
    if not matches:
        return
        
    for start_idx, end_idx in reversed(matches):
        run_indices = []
        current_len = 0
        for i, run in enumerate(paragraph.runs):
            run_len = len(run.text)
            run_start = current_len
            run_end = current_len + run_len
            if not (run_end <= start_idx or run_start >= end_idx):
                run_indices.append(i)
            current_len = run_end
            
        if not run_indices:
            continue
            
        first_run_idx = run_indices[0]
        first_run = paragraph.runs[first_run_idx]
        combined_text = "".join(paragraph.runs[idx].text for idx in run_indices)
        
        if is_whole_word:
            replaced_text = re.sub(r'\b' + re.escape(term) + r'\b', placeholder, combined_text)
        else:
            replaced_text = combined_text.replace(term, placeholder)
            
        first_run.text = replaced_text
        for idx in run_indices[1:]:
            paragraph.runs[idx].text = ""

def replace_in_paragraph(paragraph):
    pt = paragraph.text
    if not pt:
        return

    # Check for PII presence
    has_predefined = bool(union_pattern.search(pt))
    has_regex = any(c in pt for c in ["@", "+", "0", "1", "2", "3", "4", "5", "6", "7", "8", "9"])

    if not has_predefined and not has_regex:
        return

    # 1. Placeholders pass
    if has_predefined:
        for term in sorted_terms:
            if term not in pt:
                continue
            placeholder = term_to_placeholder[term]
            is_whole_word = False
            if term.isalpha() and len(term) <= 12 and term not in substring_allowed:
                is_whole_word = True
                
            replace_term_in_paragraph(paragraph, term, placeholder, is_whole_word)
        
    # Apply regex dynamic placeholders to all runs
    if has_regex:
        for run in paragraph.runs:
            if run.text:
                run.text = apply_regex_placeholders(run.text)

    # 2. Fill placeholders pass (only if run contains placeholder symbol ##)
    for run in paragraph.runs:
        if run.text and "##" in run.text:
            run.text = fill_placeholders(run.text)

def check_xml_has_pii(xml_text):
    # Quick regex checks on XML text to bypass tables/rows/cells without PII
    if union_pattern.search(xml_text):
        return True
    if any(c in xml_text for c in ["@", "+", "0", "1", "2", "3", "4", "5", "6", "7", "8", "9"]):
        return True
    return False

def replace_in_table(table):
    # Check table level
    if not check_xml_has_pii(table._element.xml):
        return
        
    for row in table.rows:
        # Check row level
        if not check_xml_has_pii(row._tr.xml):
            continue
            
        for cell in row.cells:
            # Check cell level
            if not check_xml_has_pii(cell._tc.xml):
                continue
                
            for p in cell.paragraphs:
                replace_in_paragraph(p)
            for nest_t in cell.tables:
                replace_in_table(nest_t)

def main():
    print(f"Loading document: {input_docx}")
    doc = docx.Document(input_docx)
    
    print("Redacting paragraphs...")
    for p in doc.paragraphs:
        replace_in_paragraph(p)
        
    print("Redacting tables (with XML bypassing)...")
    for t in doc.tables:
        replace_in_table(t)
        
    print("Redacting headers & footers...")
    for section in doc.sections:
        for h in [section.header, section.first_page_header, section.even_page_header]:
            if h:
                # Check header level
                if check_xml_has_pii(h._element.xml):
                    for p in h.paragraphs:
                        replace_in_paragraph(p)
                    for t in h.tables:
                        replace_in_table(t)
        for f in [section.footer, section.first_page_footer, section.even_page_footer]:
            if f:
                # Check footer level
                if check_xml_has_pii(f._element.xml):
                    for p in f.paragraphs:
                        replace_in_paragraph(p)
                    for t in f.tables:
                        replace_in_table(t)

    temp_docx = output_docx + ".temp"
    print(f"Saving text-redacted docx to: {temp_docx}")
    doc.save(temp_docx)
    
    print("Overwriting identity card media images...")
    media_replacements = {
        "word/media/image4.png": os.path.join(scratch_dir, "image4_redacted.png"),
        "word/media/image5.png": os.path.join(scratch_dir, "image5_redacted.png")
    }
    
    with zipfile.ZipFile(temp_docx, 'r') as yin:
        with zipfile.ZipFile(output_docx, 'w', zipfile.ZIP_DEFLATED) as yout:
            for item in yin.infolist():
                if item.filename in media_replacements:
                    yout.write(media_replacements[item.filename], item.filename)
                else:
                    yout.writestr(item, yin.read(item.filename))
                    
    if os.path.exists(temp_docx):
        os.remove(temp_docx)
        
    print(f"PII Redaction complete! Output saved to: {output_docx}")

if __name__ == "__main__":
    main()
