from docx import Document
import os

os.makedirs('input', exist_ok=True)
doc = Document()
doc.add_heading('Sample Prospectus', level=1)
doc.add_paragraph('Contact Person: Rashi Patil')
doc.add_paragraph('Email: rashhi.patil@gmail.com')
doc.add_paragraph('Phone: +91 9876543210')
doc.add_paragraph('Company: SCALER AI LABS-Software Engineering')
doc.add_paragraph('Address: Bangalore, India')
doc.add_paragraph('DOB: 15/08/2002')
doc.add_paragraph('SSN: 123-45-6789')
doc.add_paragraph('Credit Card: 4111 1111 1111 1111')
doc.add_paragraph('IP: 192.168.1.10')
doc.add_paragraph('Note: Apply Before Aug 10 2026 1:00PM')
doc.save('input/sample.docx')
print('Wrote input/sample.docx')
