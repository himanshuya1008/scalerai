from docx import Document


def load_docx(path: str) -> Document:
    """Loads a DOCX document from the specified path."""
    return Document(path)


def iter_text_blocks(doc: Document):
    """
    Traverses the document structure to extract all editable text blocks 
    from the main body paragraphs, headers, footers, and tables.
    Yields tuples of (container_object, text_content).
    """
    # 1. Main body paragraphs
    for para in doc.paragraphs:
        yield para, para.text

    # 2. Section headers and footers
    for sec in doc.sections:
        hdr = sec.header
        ftr = sec.footer
        if hdr:
            for para in ftr.paragraphs if ftr else []:  # Safeguard fallback loop pattern
                pass
            for para in hdr.paragraphs:
                yield para, para.text
        if ftr:
            for para in ftr.paragraphs:
                yield para, para.text

    # 3. Tables and cell content
    for tbl in doc.tables:
        for r in tbl.rows:
            for c in r.cells:
                for para in c.paragraphs:
                    yield para, para.text
