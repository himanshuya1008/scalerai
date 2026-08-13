import os
from docx import Document
from src.main import collect_detections, build_replacements_for_container
from src.anonymizer import Anonymizer
from src.writer import apply_replacements_to_container
from src.extract import iter_text_blocks
from src.detectors import detect_all


def test_all_scenarios():
    doc = Document()
    
    # 1. Full name, 2. Email, 3. Indian phone, 4. DOB, 5. Address, 6. Company, 7. Contact person, 8. SSN, 9. Credit card, 10. IPv4
    # 11. Normal order number, 12. Normal date
    # 13. Multiple occurrences of same person
    # 15. Multiple PII types in the same sentence
    doc.add_paragraph("Customer Full Name: Rashi Patil, email: rashi.patil@example.com, and phone is +91 9876543210.")
    doc.add_paragraph("Her Date of Birth is 15/08/2002. She works at Sunrise Technologies Private Limited.")
    doc.add_paragraph("Our Contact Person: Rohan Mehta can be reached at rohan.mehta@example.com.")
    doc.add_paragraph("Sensitive Data: SSN is 123-45-6789 and Credit Card is 4111 1111 1111 1111. IP Address is 192.168.1.100.")
    doc.add_paragraph("Normal details: Order number ORD-2026-1045 was placed on 20/07/2026. This is not PII.")
    doc.add_paragraph("We must ensure Rashi Patil and Rohan Mehta are replaced consistently. So Rashi Patil should map to the same name.")
    doc.add_paragraph("Address: Flat 402, Green Residency, Baner Road, Pune - 411045, Maharashtra, India")
    
    # 14. PII inside a table
    table = doc.add_table(rows=2, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Header'
    hdr_cells[1].text = 'Detail'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Contact Email'
    row_cells[1].text = 'rohan.mehta@example.com'
    
    # Save test doc
    test_path = "input/test_pii_scenarios.docx"
    doc.save(test_path)
    
    # Run redaction pipeline
    anonymizer = Anonymizer(persist_path="output/test_mapping.json")
    
    detections = []
    for container, text in iter_text_blocks(doc):
        dets = detect_all(text)
        for d in dets:
            d["container_obj"] = container
            d["container_text"] = text
        detections.extend(dets)
        
    container_reps = build_replacements_for_container(detections, anonymizer)
    for item in container_reps:
        apply_replacements_to_container(item["container"], item["replacements"])
        
    redacted_path = "output/test_pii_scenarios_redacted.docx"
    doc.save(redacted_path)
    
    # Let's inspect findings
    types = [d["type"] for d in detections]
    texts = [d["text"] for d in detections]
    
    print("Detected types:", set(types))
    print("Detections:", [f"{d['type']}: {d['text']}" for d in detections])
    
    # Assertions
    assert "PERSON" in types
    assert "EMAIL" in types
    assert "PHONE" in types
    assert "COMPANY" in types
    assert "ADDRESS" in types
    assert "SSN" in types
    assert "CREDIT_CARD" in types
    assert "DOB" in types
    assert "IP_ADDRESS" in types
    
    # Make sure labels are NOT in detected texts
    for txt in texts:
        assert not txt.endswith(":")
        assert "Address:" not in txt
        assert "Company:" not in txt
        assert "IP Address:" not in txt
        assert "Credit Card:" not in txt
        assert "Phone:" not in txt
        
    # Make sure normal order numbers and normal dates are NOT detected
    assert "ORD-2026-1045" not in texts
    assert "20/07/2026" not in texts
    
    # Verify consistent mapping
    rashi_replacements = [r["replacement"] for item in container_reps for r in item["replacements"] if r["original"] == "Rashi Patil"]
    if len(rashi_replacements) > 1:
        assert len(set(rashi_replacements)) == 1

    # Clean up files
    if os.path.exists(test_path):
        os.remove(test_path)
